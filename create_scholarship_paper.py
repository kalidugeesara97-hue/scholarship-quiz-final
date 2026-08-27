"""
Grade 5 Scholarship Exam 2026 - Guess Paper Generator
Creates a full exam paper in Sinhala medium following the official format:
  Paper I  : බුද්ධි පරීක්ෂණය (General Intelligence & Reasoning) - 40 MCQs - 1 hour
  Paper II : විෂය දැනුම (Subject Knowledge) - 1 hr 15 min
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Global styles ──────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Iskoola Pota'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

# Set default font for the document
rFonts = doc.styles['Normal'].element.rPr
if rFonts is None:
    rPr = doc.styles['Normal'].element.makeelement(qn('w:rPr'), {})
    doc.styles['Normal'].element.append(rPr)
    rFonts = rPr

# ── Helper functions ──────────────────────────────────────────────────

def add_header(text, level=1, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Iskoola Pota'
    if size:
        run.font.size = Pt(size)
    elif level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    elif level == 3:
        run.font.size = Pt(12)
    return p

def add_text(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Iskoola Pota'
    run.font.size = Pt(size)
    return p

def add_question(number, text, options=None, size=11):
    """Add an MCQ question with options."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{number}.  {text}")
    run.font.name = 'Iskoola Pota'
    run.font.size = Pt(size)
    run.bold = True

    if options:
        for i, opt in enumerate(options):
            label = chr(ord('A') + i)  # A, B, C, D
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.5)
            p2.paragraph_format.space_before = Pt(1)
            p2.paragraph_format.space_after = Pt(1)
            run2 = p2.add_run(f"  ({label})  {opt}")
            run2.font.name = 'Iskoola Pota'
            run2.font.size = Pt(size)

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

def add_section_title(text):
    add_separator()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Iskoola Pota'
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 51, 102)
    add_separator()

# ══════════════════════════════════════════════════════════════════════
#                        COVER PAGE
# ══════════════════════════════════════════════════════════════════════

for _ in range(3):
    doc.add_paragraph()

add_header("ශ්‍රී ලංකා විභාග දෙපාර්තමේන්තුව", level=2, size=14)
add_header("DEPARTMENT OF EXAMINATIONS - SRI LANKA", level=2, size=11)

doc.add_paragraph()

add_header("5 ශ්‍රේණිය ශිෂ්‍යත්ව විභාගය - 2026", level=1, size=18)
add_header("GRADE 5 SCHOLARSHIP EXAMINATION - 2026", level=1, size=14)

doc.add_paragraph()

add_header("ආදර්ශ / අනුමාන ප්‍රශ්න පත්‍රය", level=1, size=16)
add_header("(MODEL / GUESS PAPER)", level=2, size=12)

doc.add_paragraph()

add_header("සිංහල මාධ්‍යය", level=2, size=14)
add_header("SINHALA MEDIUM", level=2, size=12)

doc.add_paragraph()
doc.add_paragraph()

add_text("මෙම ප්‍රශ්න පත්‍රය පසුගිය වසරවල විභාග ප්‍රශ්න රටා විශ්ලේෂණය කර, "
         "2026 විභාගය සඳහා සකස් කරන ලද අනුමාන ප්‍රශ්න පත්‍රයකි.",
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

add_text("පත්‍ර දෙකෙහිම ප්‍රශ්න ඇතුළත් වේ.", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

doc.add_paragraph()

# Info box
info_items = [
    "I වන ප්‍රශ්න පත්‍රය : බුද්ධි පරීක්ෂණය — බහු වරණ ප්‍රශ්න 40 ක් — පැය 01 යි",
    "II වන ප්‍රශ්න පත්‍රය : විෂය දැනුම — බහු වරණ ප්‍රශ්න 60 ක් — පැය 01 යි විනාඩි 15 යි",
    "සම්පූර්ණ ලකුණු : 200"
]
for item in info_items:
    add_text(f"  ● {item}", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#              PAPER I — බුද්ධි පරීක්ෂණය
# ══════════════════════════════════════════════════════════════════════

add_header("I වන ප්‍රශ්න පත්‍රය", level=1, size=16)
add_header("බුද්ධි පරීක්ෂණය (General Intelligence & Reasoning)", level=2, size=13)
add_text("කාලය : පැය 01 යි          ලකුණු : 100          ප්‍රශ්න සංඛ්‍යාව : 40",
         bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text("සෑම ප්‍රශ්නයකටම පිළිතුරු 4 ක් දී ඇත. නිවැරදි පිළිතුර තෝරා එහි අකුර ලියන්න.",
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

doc.add_paragraph()

# ── Section A: අනුක්‍රමණ / Sequences ───────────────────────────
add_section_title("කොටස A - අනුක්‍රමණ සහ රටා හඳුනා ගැනීම")

add_question(1, "පහත සංඛ්‍යා අනුක්‍රමණයේ හිස්තැන පුරවන්න:  2, 6, 18, 54, ____",
    ["108", "162", "148", "216"])

add_question(2, "පහත අකුරු අනුක්‍රමණයේ ඊළඟ අකුර කුමක්ද?  A, C, F, J, ____",
    ["O", "N", "M", "P"])

add_question(3, "පහත සංඛ්‍යා අනුක්‍රමණයේ '?' ස්ථානයට එන සංඛ්‍යාව කුමක්ද?  5, 10, 20, 40, ?",
    ["60", "70", "80", "100"])

add_question(4, "හිස්තැනට සුදුසු සංඛ්‍යාව කුමක්ද?  3, 9, 27, ____, 243",
    ["54", "81", "72", "63"])

add_question(5, "පහත අනුක්‍රමණය අධ්‍යයනය කරන්න: 1, 4, 9, 16, 25, ____. හිස්තැනට එන සංඛ්‍යාව කුමක්ද?",
    ["30", "36", "49", "32"])

# ── Section B: සාමාන්‍ය බුද්ධිය / Verbal Reasoning ─────────
add_section_title("කොටස B - වාචික තර්කනය")

add_question(6, "'පොත' යන්නට 'කියවීම' වැනිද, 'ආහාර' යන්නට ____ වැනිය.",
    ["රසය", "කෑම", "උයනය", "ගෙවත්ත"])

add_question(7, "පහත කුමණ වචනය අනිත් වචන වලට වෙනස්ද?",
    ["ගස", "මල", "කොළ", "ගල"])

add_question(8, "'ඉර' යන්නට 'දිවා' වැනිද, 'හඳ' යන්නට ____ වැනිය.",
    ["රාත්‍රිය", "අඳුර", "තරු", "වළාකුළ"])

add_question(9, "A යනු B ගේ සහෝදරයාය. B යනු C ගේ මවය. එසේ නම් A යනු C ට කුමක්ද?",
    ["මාමා", "පියා", "සහෝදරයා", "මුත්තා"])

add_question(10, "පහත වචන අතරින් අනිත් වචනවලට අයත් නොවන එක කුමක්ද? — බස්, ට්‍රේන්, බයිසිකලය, පුටුව",
    ["බස්", "ට්‍රේන්", "බයිසිකලය", "පුටුව"])

# ── Section C: ගණිත තාර්කිකත්වය ─────────────────────────
add_section_title("කොටස C - ගණිත තාර්කිකත්වය")

add_question(11, "යම් සංඛ්‍යාවකට 15 එකතු කළ විට 42 ලැබේ නම්, එම සංඛ්‍යාව කුමක්ද?",
    ["25", "27", "32", "37"])

add_question(12, "රුපියල් 500 ක මුදලකින් රුපියල් 175 ක භාණ්ඩයක් මිලදී ගත් විට ඉතිරි මුදල කොපමණද?",
    ["රු. 225", "රු. 325", "රු. 375", "රු. 275"])

add_question(13, "එක් පෙට්ටියක බෝල 12 ක් ඇත. එවැනි පෙට්ටි 8 ක ඇති බෝල ගණන කීයද?",
    ["80", "96", "84", "108"])

add_question(14, "නිමල්ගේ වයස අවුරුදු 9 කි. ඔහුගේ පියාගේ වයස නිමල්ගේ වයසට 4 ගුණයක් නම්, පියාගේ වයස කීයද?",
    ["32", "36", "40", "45"])

add_question(15, "පන්තියක සිසුන් 35 දෙනෙකු සිටිති. ඉන් 3/5 ක් ගැහැණු සිසුවියන් නම්, පිරිමි සිසුන් ගණන කීයද?",
    ["21", "14", "15", "20"])

add_question(16, "දිනයක පැය 24 කි. දිනයකින් 3/8 ක් ගත වී ඇත නම්, ගත වූ පැය ගණන කීයද?",
    ["6", "8", "9", "12"])

add_question(17, "ත්‍රිකෝණයක කෝණ තුන 60°, 70° සහ ____ නම්, හිස්තැනට එන අගය කුමක්ද?",
    ["40°", "50°", "60°", "70°"])

add_question(18, "සෘජුකෝණාස්‍රයක දිග 8 cm ද, පළල 5 cm ද වේ. එහි වර්ගඵලය කුමක්ද?",
    ["26 cm²", "40 cm²", "13 cm²", "80 cm²"])

# ── Section D: රූප හඳුනා ගැනීම / Non-verbal ─────────────
add_section_title("කොටස D - අවාචික තර්කනය / රූප හඳුනාගැනීම")

add_question(19, "පහත රූප අනුක්‍රමණයේ ඊළඟ රූපය කුමක්ද? ○ □ △ ○ □ ____",
    ["○", "□", "△", "◇"])

add_question(20, "කැඩපතකින් බැලූ විට 'b' අකුර පෙනෙන්නේ කෙසේද?",
    ["d", "p", "q", "b"])

add_question(21, "පහත සරල රටාවෙහි '?' ස්ථානයට සුදුසු වන්නේ කුමක්ද?\n    1 → 1\n    2 → 4\n    3 → 9\n    4 → ?",
    ["12", "16", "14", "8"])

add_question(22, "පහත රූප කණ්ඩායමෙන් වෙනස් වූ එක තෝරන්න: ▲ ▲ ▼ ▲",
    ["පළමු", "දෙවැනි", "තුන්වැනි", "සිව්වැනි"])

add_question(23, "සතරැස් කඩදාසියක් දෙකට නැම්මූ විට ලැබෙන හැඩය කුමක්ද?",
    ["ත්‍රිකෝණය", "සෘජුකෝණාස්‍රය", "වෘත්තය", "ෂඩ්භුජය"])

# ── Section E: කේත සහ තර්කනය ─────────────────────────────
add_section_title("කොටස E - කේත භාෂාව සහ තර්කනය")

add_question(24, "'CAT' යන වචනය කේත භාෂාවකින් 'DBU' ලෙස ලියනු ලැබේ නම් 'DOG' කේතනය කුමක්ද?",
    ["EPH", "CPF", "FQI", "EOG"])

add_question(25, "A > B, B > C, C > D නම් විශාලම අකුර කුමක්ද?",
    ["A", "B", "C", "D"])

add_question(26, "යම් මාසයක 1 වැනිදා සඳුදා නම්, එම මාසයේ 15 වැනිදා කුමන දිනයකට වැටේද?",
    ["සඳුදා", "බදාදා", "සිකුරාදා", "ඉරිදා"])

add_question(27, "දකුණට හැරී සිටින පුද්ගලයෙකු වමට අංශක 90 ක් හැරුණහොත් ඔහු බලන දිශාව කුමක්ද?",
    ["උතුර", "දකුණ", "නැගෙනහිර", "බටහිර"])

# ── Section F: සාමාන්‍ය දැනුම / Comprehension ──────────────
add_section_title("කොටස F - කෙටි ගද්‍ය හා තේරුම් ගැනීම")

add_text("පහත කෙටි ගද්‍යය කියවා ප්‍රශ්නවලට පිළිතුරු දෙන්න:", bold=True, size=11)
add_text('"නිමල් සෑම දිනකම උදේ 6.00 ට පිබිදේ. ඔහු මුහුණ සෝදා ආහාර ගෙන පාසලට යයි. '
         'ඔහුගේ පාසල ගෙදරින් කිලෝමීටර 2 ක් දුරින් පිහිටා ඇත. ඔහු බස් රථයෙන් පාසලට යයි. '
         'පාසලේ ඔහුගේ හොඳම මිතුරා සුනිල්ය. ඔවුන් දෙදෙනා එකට සෙල්ලම් කරති."',
         size=10, italic=True)

add_question(28, "නිමල් උදේ පිබිදෙන්නේ කීයටද?",
    ["5.00 ට", "6.00 ට", "7.00 ට", "6.30 ට"])

add_question(29, "නිමල් පාසලට යන්නේ කෙසේද?",
    ["පයින්", "බයිසිකලයෙන්", "බස් රථයෙන්", "වෑන් රථයෙන්"])

add_question(30, "නිමල්ගේ හොඳම මිතුරා කවුද?",
    ["කමල්", "සුනිල්", "අමල්", "නිශාන්ත"])

# ── Section G: තවත් තර්කන ප්‍රශ්න ────────────────────────
add_section_title("කොටස G - මිශ්‍ර තර්කන ප්‍රශ්න")

add_question(31, "ඔරලෝසුවක කටු 3.00 පෙන්වන විට, පැය කටුව සහ මිනිත්තු කටුව අතර කෝණය කුමක්ද?",
    ["60°", "90°", "120°", "180°"])

add_question(32, "5 දෙනෙකු වෘත්තාකාරව වාඩි වී සිටිති. A, B ට දකුණෙන් සිටී. C, B ට වමෙන් සිටී. D, C ට දකුණෙන් සිටී. E, A ට වමෙන් සිටී. B ට මුහුණලා සිටින්නේ කවුද?",
    ["A", "D", "E", "C"])

add_question(33, "පහත කුමණ සංඛ්‍යාව 3 නුත් 5 නුත් වලින් බෙදිය හැකිද?",
    ["10", "12", "15", "20"])

add_question(34, "කාලය දැන් උදේ 9.45 නම්, මිනිත්තු 50 කට පසුව කාලය කුමක්ද?",
    ["10.25", "10.35", "10.45", "10.55"])

add_question(35, "සෘජුකෝණාස්‍රාකාර තට්ටුවක දිග 12 m, පළල 8 m නම් එහි පරිමිතිය කුමක්ද?",
    ["40 m", "96 m", "20 m", "36 m"])

add_question(36, "සතියක දින 7 කි. සති 4 ක සහ දින 3 ක එකතුව කොපමණද?",
    ["28 දින", "30 දින", "31 දින", "35 දින"])

add_question(37, "පුස්තකාලයක පොත් 856 ක් ඇත. ඉන් 289 ක් සිසුන්ට ලබා දී ඇත. ඉතිරි පොත් ගණන කීයද?",
    ["567", "577", "547", "667"])

add_question(38, "ගමනක් ආරම්භ කළ වාහනයක මීටර් කියවීම 45123 කි. ගමන අවසානයේ මීටර් කියවීම 45398 නම්, ගිය දුර කුමක්ද?",
    ["275 km", "375 km", "225 km", "265 km"])

add_question(39, "එක් කේක් එකක මිල රු.45 නම්, කේක් 7 ක මිල කීයද?",
    ["රු. 285", "රු. 315", "රු. 305", "රු. 295"])

add_question(40, "අවුරුද්දක මාස 12 කි. මාස 30 ක් යනු අවුරුදු කීයක් සහ මාස කීයක්ද?",
    ["2 අවුරුදු 6 මාස", "3 අවුරුදු", "2 අවුරුදු 8 මාස", "1 අවුරුද්ද 6 මාස"])

add_text("——— I වන ප්‍රශ්න පත්‍රයේ අවසානයයි ———",
         bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#     PAPER II — විෂය දැනුම (Subject Knowledge)
# ══════════════════════════════════════════════════════════════════════

add_header("II වන ප්‍රශ්න පත්‍රය", level=1, size=16)
add_header("විෂය දැනුම (Subject Knowledge)", level=2, size=13)
add_text("කාලය : පැය 01 යි විනාඩි 15 යි     ලකුණු : 100     ප්‍රශ්න සංඛ්‍යාව : 60",
         bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text("සෑම ප්‍රශ්නයකටම පිළිතුරු 4 ක් දී ඇත. නිවැරදි පිළිතුර තෝරා එහි අකුර ලියන්න.",
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

doc.add_paragraph()

# ═══════════════════════════════════════════
#  SECTION 1: මව්බස — සිංහල (Q 1-15)
# ═══════════════════════════════════════════
add_section_title("කොටස 1 — මව්බස (සිංහල) — ප්‍රශ්න 01 - 15")

add_question(1, "පහත වාක්‍යයේ නිවැරදි අකුරු වින්‍යාසය සහිත වචනය කුමක්ද?\n    'ළමයා _______ කියවයි.'",
    ["පොත", "පෝත", "පෝද", "බොත"])

add_question(2, "'ගුරුවරයා' යන වචනයේ ස්ත්‍රී ලිංගය කුමක්ද?",
    ["ගුරුවරිය", "ගුරුවර", "ගුරුතුමිය", "ගුරුතුමා"])

add_question(3, "පහත වාක්‍යයේ හිස්තැනට සුදුසු වචනය තෝරන්න:\n    'අම්මා _______ උයනවා.'",
    ["බත්", "බත", "බැත්", "බැත"])

add_question(4, "'ඉගෙනීම' යන වචනයේ විරුද්ධාර්ථ වචනය කුමක්ද?",
    ["ඉගැන්වීම", "නොඉගෙනීම", "දැනීම", "කියවීම"])

add_question(5, "පහත කුමණ වාක්‍යය ව්‍යාකරණ වශයෙන් නිවැරදිද?",
    ["මම පාසලට යනවා.", "මම පාසලට යනව.", "මම පාසල යනවා.", "මම පාසලටම යනව."])

add_question(6, "'සූර්යයා' යන වචනයේ සමානාර්ථ වචනයක් වන්නේ කුමක්ද?",
    ["ඉර", "හඳ", "තරුව", "අහස"])

add_question(7, "'ගස' යන වචනයේ බහු වචනය කුමක්ද?",
    ["ගස්", "ගස්වල", "ගස්හි", "ගසක්"])

add_question(8, "පහත පද්‍යයේ හිස්තැනට සුදුසු වචනය කුමක්ද?\n    'කුරුල්ලෝ ____ ගයති'",
    ["සිංදු", "ගී", "නාද", "ගීත"])

add_question(9, "'දරුවන්ට' යන වචනයේ ඇති විභක්ති ප්‍රත්‍යය කුමක්ද?",
    ["-ට", "-ගේ", "-ගෙන්", "-ව"])

add_question(10, "පහත කුමණ වචනය තත්සම පදයක්ද?",
    ["පුෂ්පය", "මල", "හුරතලයා", "දරුවා"])

add_question(11, "'ගුවන්විදුලිය' යන සංයුක්ත පදය බිඳ දක්වන්න:",
    ["ගුවන් + විදුලිය", "ගුවන්වි + දුලිය", "ගුව + න්විදුලිය", "ගුවන්විදු + ලිය"])

add_question(12, "පහත වාක්‍යයේ යටින් ඉරි ඇඳ ඇති පදයේ කර්ම කුමක්ද?\n    'ළමයා පොත කියවයි.'",
    ["පොත", "ළමයා", "කියවයි", "යි"])

add_question(13, "'ආරක්ෂාව' යන වචනයේ මූල ක්‍රියා පදය කුමක්ද?",
    ["ආරක්ෂා කිරීම", "ආරක්ෂිත", "රක්ෂණය", "ආරක්ෂක"])

add_question(14, "පහත කුමණ වචනය නිවැරදිව ලියා ඇත්ද?",
    ["විද්‍යාව", "විද්යාව", "විධ්‍යාව", "විද්ධාව"])

add_question(15, "'අම්මා කෑම හදනවා' යන වාක්‍යයේ කර්තෘ පදය කුමක්ද?",
    ["අම්මා", "කෑම", "හදනවා", "කෑම හදනවා"])

# ═══════════════════════════════════════════
#  SECTION 2: ගණිතය (Q 16-35)
# ═══════════════════════════════════════════
add_section_title("කොටස 2 — ගණිතය — ප්‍රශ්න 16 - 35")

add_question(16, "3456 + 2789 = ?",
    ["6245", "6345", "6145", "6235"])

add_question(17, "8000 - 3567 = ?",
    ["4433", "4333", "4533", "4343"])

add_question(18, "256 × 4 = ?",
    ["1024", "1124", "924", "1034"])

add_question(19, "7296 ÷ 8 = ?",
    ["912", "812", "902", "892"])

add_question(20, "1/4 + 2/4 = ?",
    ["3/8", "3/4", "1/2", "2/4"])

add_question(21, "4.5 + 3.25 = ?",
    ["7.75", "7.25", "8.75", "7.70"])

add_question(22, "පහත කුමණ භාගය විශාලම භාගයද?",
    ["1/2", "1/3", "1/4", "1/5"])

add_question(23, "සෘජුකෝණාස්‍රයක දිග 15 cm, පළල 10 cm නම්, වර්ගඵලය කුමක්ද?",
    ["150 cm²", "50 cm²", "25 cm²", "250 cm²"])

add_question(24, "අඟල් 12 ක් = ?",
    ["1 අඩිය", "2 අඩි", "1/2 අඩිය", "3 අඩි"])

add_question(25, "කිලෝග්‍රෑම් 2 ක් = ග්‍රෑම් කීයද?",
    ["200", "2000", "20000", "20"])

add_question(26, "වෘත්තයක අරය 7 cm නම්, එහි විෂ්කම්භය කුමක්ද?",
    ["14 cm", "21 cm", "7 cm", "28 cm"])

add_question(27, "පැය 2 යි විනාඩි 30 = විනාඩි කීයද?",
    ["130", "150", "120", "230"])

add_question(28, "සමචතුරස්‍රයක එක් පැත්තක දිග 9 cm නම්, එහි පරිමිතිය කුමක්ද?",
    ["27 cm", "36 cm", "81 cm", "18 cm"])

add_question(29, "475 රුපියල් 25 ශතවලට සමාන මුදල් කාසි (රු.25 බැගින්) ගණන කීයද?",
    ["19", "18", "20", "17"])

add_question(30, "3/5 දශම භාගයක් ලෙස ලියනු ලබන්නේ කෙසේද?",
    ["0.6", "0.35", "0.53", "0.06"])

add_question(31, "ත්‍රිකෝණයක පාදය 10 cm, උස 6 cm නම් වර්ගඵලය කුමක්ද?",
    ["60 cm²", "30 cm²", "16 cm²", "40 cm²"])

add_question(32, "රෝමානු ඉලක්කම්වලින් 49 ලියන්නේ කෙසේද?",
    ["XXXXIX", "XLIX", "IL", "XLVIIII"])

add_question(33, "බෝතලයක ජලය ලීටර් 1.5 කි. එවැනි බෝතල් 4 ක ඇති ජලය කුමක්ද?",
    ["4.5 l", "6 l", "5 l", "4 l"])

add_question(34, "දිනයක පැය 24 කි. දින 3 ක සහ පැය 6 ක එකතුව පැය කීයද?",
    ["72", "78", "84", "70"])

add_question(35, "පහත කුමණ සංඛ්‍යාව ප්‍රථමක සංඛ්‍යාවක්ද?",
    ["11", "9", "15", "21"])

# ═══════════════════════════════════════════
#  SECTION 3: පරිසරය (Q 36-50)
# ═══════════════════════════════════════════
add_section_title("කොටස 3 — පරිසරය ආශ්‍රිත ක්‍රියාකාරකම් — ප්‍රශ්න 36 - 50")

add_question(36, "ශ්‍රී ලංකාවේ දිගම ගඟ කුමක්ද?",
    ["මහවැලි ගඟ", "කැලණි ගඟ", "වලවේ ගඟ", "කළු ගඟ"])

add_question(37, "පැළෑටියක ආහාර නිපදවීමේ ක්‍රියාවලිය හඳුන්වන්නේ කුමක් ලෙසද?",
    ["ප්‍රකාෂ සංස්ලේෂණය", "ශ්වසනය", "පිටාසාරය", "ප්‍රජනනය"])

add_question(38, "ශ්‍රී ලංකාවේ ජාතික මල කුමක්ද?",
    ["නිල් මාහනෙල්", "රෝස මල", "සූරියකාන්ත", "ඕකිඩ්"])

add_question(39, "ජලය සිදුරු සහිත බඳුනකින් බිංදු වශයෙන් පෙරීම නිදසුනක් වන්නේ කුමක් සඳහාද?",
    ["පෙරීම", "වාෂ්පීකරණය", "ඝනීභවනය", "ද්‍රාවණය"])

add_question(40, "පහත කුමණ සතෙකුට කොඳු ඇට පෙළක් නැත?",
    ["බල්ලා", "මාළුවා", "සමනලයා", "කිරිල්ලා"])

add_question(41, "ශ්‍රී ලංකාවේ පළාත් ගණන කීයද?",
    ["7", "8", "9", "10"])

add_question(42, "පෘථිවිය වටා භ්‍රමණය වන ග්‍රහයා කුමක්ද? (නැත, පෘථිවිය වටා කුමක් භ්‍රමණය වේද?)",
    ["සඳ", "ඉරි", "අඟහරු", "සිකුරු"])

add_question(43, "ගස්වලින් පරිසරයට මුදා හරින වායුව කුමක්ද?",
    ["ඔක්සිජන්", "කාබන් ඩයොක්සයිඩ්", "නයිට්‍රජන්", "හයිඩ්‍රජන්"])

add_question(44, "ශ්‍රී ලංකාවේ ජාතික ගසෙහි නම කුමක්ද?",
    ["නා ගස", "කොස් ගස", "පොල් ගස", "මාර ගස"])

add_question(45, "පහත කුමණ ආහාරය ශරීරයට බලශක්ති ලබා දේද?",
    ["සහල්", "ජලය", "ලුණු", "විනාකිරි"])

add_question(46, "ශ්‍රී ලංකාවේ අගනගරය කුමක්ද?",
    ["කොළඹ", "ශ්‍රී ජයවර්ධනපුර කෝට්ටේ", "මහනුවර", "ගාල්ල"])

add_question(47, "වර්ෂාව මැනීම සඳහා භාවිතා කරන උපකරණය කුමක්ද?",
    ["වර්ෂා මානය", "උෂ්ණත්ව මානය", "බැරෝමීටරය", "තුලාව"])

add_question(48, "පෘථිවියේ මතුපිට වැඩි ප්‍රමාණයක් ආවරණය වී ඇත්තේ කුමකින්ද?",
    ["ජලයෙන්", "පසින්", "වනාන්තරවලින්", "හිමෙන්"])

add_question(49, "ශ්‍රී ලංකාවේ උසම කන්ද කුමක්ද?",
    ["පිදුරුතලාගල", "ඇලාගල", "ලක්ෂපාන", "බයිබල් කන්ද"])

add_question(50, "ජලය උනුසුම් කළ විට සිදුවන ක්‍රියාවලිය කුමක්ද?",
    ["වාෂ්පීකරණය", "ඝනීභවනය", "පෙරීම", "ද්‍රවණය"])

# ═══════════════════════════════════════════
#  SECTION 4: ආගම හා සදාචාරය (Q 51-55)
# ═══════════════════════════════════════════
add_section_title("කොටස 4 — බුද්ධ ධර්මය / ආගම සහ සදාචාරය — ප්‍රශ්න 51 - 55")

add_question(51, "බුදුරජාණන් වහන්සේ උපන් ස්ථානය කුමක්ද?",
    ["ලුම්බිනිය", "බුද්ධගයාව", "සාරානාත්", "කුසිනාරාව"])

add_question(52, "පංච ශීලයේ පළමු ශීලය කුමක්ද?",
    ["ප්‍රාණඝාතයෙන් වැළකීම", "අදත්තාදානයෙන් වැළකීම", "මුසාවාදයෙන් වැළකීම", "සුරාමේරයෙන් වැළකීම"])

add_question(53, "චතුරාර්ය සත්‍යයන් අතර පළමු සත්‍යය කුමක්ද?",
    ["දුක්ඛ සත්‍යය", "සමුදය සත්‍යය", "නිරෝධ සත්‍යය", "මාර්ග සත්‍යය"])

add_question(54, "බුදුරජාණන් වහන්සේ ප්‍රථම ධර්ම දේශනය පැවැත්වූ ස්ථානය කුමක්ද?",
    ["මිගදාය (සාරානාත්)", "ලුම්බිනිය", "බුද්ධගයාව", "රාජගහ"])

add_question(55, "පහත කුමණ ගුණාංගය බුද්ධ ධර්මයේ අගය කෙරේද?",
    ["මෛත්‍රිය", "කෝපය", "ලෝභය", "මසුරුකම"])

# ═══════════════════════════════════════════
#  SECTION 5: English (Q 56-60)
# ═══════════════════════════════════════════
add_section_title("කොටස 5 — English — ප්‍රශ්න 56 - 60")

add_question(56, "Choose the correct word to fill the blank:\n    'She ____ to school every day.'",
    ["goes", "go", "going", "gone"])

add_question(57, "What is the plural of 'child'?",
    ["children", "childs", "childrens", "childes"])

add_question(58, "Choose the correct sentence:",
    ["The cat is sleeping on the mat.", "The cat is sleep on the mat.", "The cat are sleeping on the mat.", "The cat is sleeps on the mat."])

add_question(59, "What is the opposite of 'big'?",
    ["small", "tall", "long", "wide"])

add_question(60, "Fill in the blank: 'I ____ a student.'",
    ["am", "is", "are", "be"])

doc.add_paragraph()
add_text("——— II වන ප්‍රශ්න පත්‍රයේ අවසානයයි ———",
         bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#              ANSWER KEY / පිළිතුරු පත්‍රය
# ══════════════════════════════════════════════════════════════════════

add_header("පිළිතුරු පත්‍රය / ANSWER KEY", level=1, size=16)

doc.add_paragraph()
add_header("I වන ප්‍රශ්න පත්‍රය — බුද්ධි පරීක්ෂණය", level=2, size=13)

paper1_answers = {
    1: "B", 2: "A", 3: "C", 4: "B", 5: "B",
    6: "B", 7: "D", 8: "A", 9: "A", 10: "D",
    11: "B", 12: "B", 13: "B", 14: "B", 15: "B",
    16: "C", 17: "B", 18: "B", 19: "C", 20: "A",
    21: "B", 22: "C", 23: "B", 24: "A", 25: "A",
    26: "A", 27: "A", 28: "B", 29: "C", 30: "B",
    31: "B", 32: "B", 33: "C", 34: "B", 35: "A",
    36: "B", 37: "A", 38: "A", 39: "B", 40: "A"
}

# Create answer table for Paper I
table1 = doc.add_table(rows=9, cols=10)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.style = 'Table Grid'

# Header row
for col_idx in range(0, 10, 2):
    q_cell = table1.cell(0, col_idx)
    q_cell.text = "ප්‍රශ්නය"
    a_cell = table1.cell(0, col_idx + 1)
    a_cell.text = "පිළිතුර"
    for cell in [q_cell, a_cell]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Iskoola Pota'

q_num = 1
for row_idx in range(1, 9):
    for col_pair in range(0, 10, 2):
        if q_num <= 40:
            table1.cell(row_idx, col_pair).text = str(q_num)
            table1.cell(row_idx, col_pair + 1).text = paper1_answers[q_num]
            for cell in [table1.cell(row_idx, col_pair), table1.cell(row_idx, col_pair + 1)]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Iskoola Pota'
            q_num += 1

doc.add_paragraph()
add_header("II වන ප්‍රශ්න පත්‍රය — විෂය දැනුම", level=2, size=13)

paper2_answers = {
    1: "A", 2: "A", 3: "A", 4: "A", 5: "A",
    6: "A", 7: "A", 8: "B", 9: "A", 10: "A",
    11: "A", 12: "A", 13: "A", 14: "A", 15: "A",
    16: "A", 17: "A", 18: "A", 19: "A", 20: "B",
    21: "A", 22: "A", 23: "A", 24: "A", 25: "B",
    26: "A", 27: "B", 28: "B", 29: "A", 30: "A",
    31: "B", 32: "B", 33: "B", 34: "B", 35: "A",
    36: "A", 37: "A", 38: "A", 39: "A", 40: "C",
    41: "C", 42: "A", 43: "A", 44: "A", 45: "A",
    46: "B", 47: "A", 48: "A", 49: "A", 50: "A",
    51: "A", 52: "A", 53: "A", 54: "A", 55: "A",
    56: "A", 57: "A", 58: "A", 59: "A", 60: "A"
}

table2 = doc.add_table(rows=13, cols=10)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Table Grid'

for col_idx in range(0, 10, 2):
    q_cell = table2.cell(0, col_idx)
    q_cell.text = "ප්‍රශ්නය"
    a_cell = table2.cell(0, col_idx + 1)
    a_cell.text = "පිළිතුර"
    for cell in [q_cell, a_cell]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Iskoola Pota'

q_num = 1
for row_idx in range(1, 13):
    for col_pair in range(0, 10, 2):
        if q_num <= 60:
            table2.cell(row_idx, col_pair).text = str(q_num)
            table2.cell(row_idx, col_pair + 1).text = paper2_answers[q_num]
            for cell in [table2.cell(row_idx, col_pair), table2.cell(row_idx, col_pair + 1)]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Iskoola Pota'
            q_num += 1

doc.add_paragraph()
add_text("★ මෙම ප්‍රශ්න පත්‍රය පසුගිය වසරවල විභාග රටා විශ්ලේෂණය මත පදනම්ව සකස් කළ අනුමාන ප්‍රශ්න පත්‍රයකි.",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text("★ මෙය නිල විභාග දෙපාර්තමේන්තුවේ ප්‍රශ්න පත්‍රයක් නොවේ.",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_text("★ වැඩිදුර සූදානම සඳහා නිල පසුගිය ප්‍රශ්න පත්‍ර සහ ආදර්ශ ප්‍රශ්න පත්‍ර භාවිතා කරන්න.",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──────────────────────────────────────────────────────────────
output_path = r"C:\Users\HP\.gemini\antigravity\scratch\Grade_5_Scholarship_Exam_2026_Guess_Paper.docx"
doc.save(output_path)
print(f"✅ Paper saved to: {output_path}")
